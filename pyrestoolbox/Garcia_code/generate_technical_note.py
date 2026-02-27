"""Generate technical note docx for Whitson review."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# Title
title = doc.add_heading('Multi-Gas Density Corrections for Aqueous Solutions', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Extending Garcia (2001) via the Plyasunov Apparent Molar Volume Model')
run.italic = True
run.font.size = Pt(13)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('Mark Burgoyne')
run.font.size = Pt(11)

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date.add_run('February 2026 — Technical Note (Draft)')
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()

# ============================================================
doc.add_heading('1. Background', level=1)

doc.add_paragraph(
    'Garcia (2001) developed a density mixing rule for aqueous solutions containing '
    'dissolved CO\u2082. His Equation 18 relates the solution density to the solvent '
    'density, the dissolved gas mole fraction, the gas molecular weight, and the '
    'apparent molar volume (V\u03c6) of the dissolved gas:'
)

eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq.add_run('\u03c1 = (1 + x\u2082\u00b7M\u2082 / (M\u2081\u00b7x\u2081)) / (x\u2082\u00b7V\u03c6 / (M\u2081\u00b7x\u2081) + 1/\u03c1\u2081)')
run.italic = True

doc.add_paragraph(
    'where \u03c1\u2081 is the solvent (water or brine) density, x\u2082 is the dissolved gas mole fraction, '
    'M\u2081 and M\u2082 are the molecular weights of water and gas respectively, and '
    'V\u03c6 is the apparent molar volume of the dissolved gas in cm\u00b3/mol.'
)

doc.add_paragraph(
    'Garcia provided a cubic polynomial for V\u03c6 of CO\u2082 as a function of temperature, '
    'but his mixing rule is entirely gas-generic. Extending it to other gases '
    'only requires providing V\u03c6 for each gas. This technical note describes how '
    'V\u03c6 can be computed for 8 dissolved gases (CO\u2082, CH\u2084, C\u2082H\u2086, C\u2083H\u2088, '
    'n-C\u2084H\u2081\u2080, H\u2082S, H\u2082, N\u2082) using the Plyasunov A\u2081\u2082\u221e model, '
    'and presents validation results.'
)

# ============================================================
doc.add_heading('2. The Plyasunov Model', level=1)

doc.add_paragraph(
    'Plyasunov and co-workers (2019-2021) developed a unified framework for computing '
    'the infinite-dilution partial molar volume V\u2082\u221e of dissolved gases in water, '
    'based on Fluctuation Solution Theory. At the dilute concentrations relevant to '
    'reservoir systems, V\u2082\u221e equals V\u03c6, making it directly substitutable into '
    'Garcia\'s mixing rule.'
)

doc.add_paragraph(
    'The model computes a dimensionless quantity A\u2081\u2082\u221e as a polynomial in pure water '
    'density, with temperature-dependent coefficients specific to each gas:'
)

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq2.add_run('A\u2081\u2082\u221e = 1 + \u03c1\u2081*\u00b7[a\u2080 + a\u2081\u00b7\u03c1\u2081* + a\u2082\u00b7(\u03c1\u2081*)\u00b2 + a\u2083\u00b7(\u03c1\u2081*)\u00b3 + a\u2084\u00b7(\u03c1\u2081*)\u2074 + a\u2085\u00b7(\u03c1\u2081*)\u2075]')
run.italic = True

doc.add_paragraph(
    'where \u03c1\u2081* is the pure water density in kg/m\u00b3 at (T, P). The apparent molar volume '
    'is then recovered via:'
)

eq3 = doc.add_paragraph()
eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq3.add_run('V\u2082\u221e(T, P) = A\u2081\u2082\u221e(T, \u03c1\u2081*) \u00b7 \u03baT(T, P) \u00b7 R \u00b7 T')
run.italic = True

doc.add_paragraph(
    'where \u03baT is the isothermal compressibility of pure water (from IAPWS-95) '
    'and R = 8.314 J/(mol\u00b7K).'
)

doc.add_paragraph(
    'The coefficient a\u2080 derives from the cross second virial coefficient B\u2081\u2082(T) '
    'between water and the solute. Coefficients a\u2081 through a\u2085 are expressed as '
    'polynomials in inverse reduced temperature (\u03b8 = T/T_c, T_c = 647.096 K) with '
    '7 fitted parameters each, giving 35 parameters per gas. Three different functional '
    'forms are used for B\u2081\u2082 depending on the gas type.'
)

doc.add_paragraph(
    'The model is published across three papers covering all 8 target gases:'
)

bullets = [
    'Part II (Plyasunov & Korzhinskaya 2020): CO\u2082, C\u2082H\u2086, C\u2083H\u2088, n-C\u2084H\u2081\u2080',
    'Part III (Plyasunov & Korzhinskaya 2021): H\u2082S',
    'Part IV (Plyasunov & Korzhinskaya 2021): H\u2082, N\u2082, CH\u2084 (supersedes Part I)',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'All three papers use the same master equation form, enabling a single unified implementation.'
)

# ============================================================
doc.add_heading('3. Mixed Dissolved Gases', level=1)

doc.add_paragraph(
    'For solutions containing multiple dissolved gases, we compute effective properties '
    'using mole-fraction weighting among the dissolved gases:'
)

eq4 = doc.add_paragraph()
eq4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = eq4.add_run('V\u03c6,eff = \u03a3 y\u1d62 \u00b7 V\u03c6,i        M\u2082,eff = \u03a3 y\u1d62 \u00b7 M\u2082,i')
run.italic = True

doc.add_paragraph(
    'where y\u1d62 = x\u2082,i / x\u2082,total is the fraction of gas i among all dissolved gases. '
    'Garcia\'s Eq. 18 is then applied with V\u03c6,eff, M\u2082,eff, and x\u2082,total. '
    'This is the standard ideal mixing approximation for apparent molar volumes, '
    'appropriate at the dilute concentrations typical of reservoir systems.'
)

# ============================================================
doc.add_heading('4. Validation', level=1)

doc.add_heading('4.1 Reference Values at 298.15 K', level=2)

doc.add_paragraph(
    'Table 1 shows the computed V\u03c6 at 298.15 K and 0.1 MPa compared to Plyasunov\'s '
    'critically evaluated reference values, along with the density effect at 2 mol% '
    'dissolved gas in pure water at 10 MPa.'
)

# Table 1
table1 = doc.add_table(rows=9, cols=6)
table1.style = 'Light Shading Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Gas', 'MW (g/mol)', 'V\u03c6 computed', 'V\u03c6 reference', '% Error', '\u0394\u03c1/\u03c1 at x\u2082=0.02']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

data = [
    ('H\u2082',      '2.016',  '26.12', '26.1', '0.08%', '\u22122.64%'),
    ('N\u2082',     '28.013', '34.71', '34.7', '0.04%', '\u22120.75%'),
    ('CH\u2084',    '16.043', '37.02', '37.0', '0.06%', '\u22122.29%'),
    ('CO\u2082',    '44.010', '34.02', '34.0', '0.06%', '+1.10%'),
    ('C\u2082H\u2086',   '30.069', '51.44', '51.4', '0.07%', '\u22122.30%'),
    ('C\u2083H\u2088',   '44.096', '67.12', '67.1', '0.03%', '\u22122.42%'),
    ('n\u2011C\u2084H\u2081\u2080', '58.122', '82.79', '82.8', '0.01%', '\u22122.57%'),
    ('H\u2082S',    '34.081', '34.81', '34.8', '0.04%', '\u22120.10%'),
]

for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table1.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph('Table 1: ')
p.runs[0].bold = True
p.runs[0].font.size = Pt(9)
run = p.add_run('V\u03c6 at 298.15 K and density effects at 2 mol% dissolved gas. '
    'CO\u2082 is the only gas that increases solution density. H\u2082S is nearly neutral.')
run.font.size = Pt(9)
run.italic = True

doc.add_heading('4.2 Temperature Dependence', level=2)

doc.add_paragraph(
    'We validated our Plyasunov model implementation by comparing computed V\u03c6 values '
    'against the experimental data of Hnedkovsky, Wood & Majer (1996, J. Chem. '
    'Thermodynamics 28, 125-142, DOI: 10.1006/jcht.1996.0011). Hnedkovsky et al. '
    'measured V\u03c6 using a vibrating tube flow densitometer at pressures of 28 and '
    '35 MPa. We compared our model predictions at 30 MPa against the average of their '
    '28 and 35 MPa measurements for CH\u2084, CO\u2082, and H\u2082S. '
    'Maximum deviations: CH\u2084 2.4%, CO\u2082 2.6%, H\u2082S 3.8%. These are within the '
    'experimental uncertainties of the densimetric measurements.'
)

# Table 2 - Temperature validation
table2 = doc.add_table(rows=7, cols=7)
table2.style = 'Light Shading Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['T (\u00b0C)', 'CH\u2084 calc', 'CH\u2084 expt', 'CO\u2082 calc', 'CO\u2082 expt', 'H\u2082S calc', 'H\u2082S expt']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

t_data = [
    ('25',  '36.92', '36.75', '33.55', '33.45', '35.05', '34.90'),
    ('50',  '37.00', '37.30', '33.48', '33.75', '35.05', '35.90'),
    ('100', '39.99', '40.50', '37.20', '37.50', '37.93', '38.70'),
    ('150', '45.62', '45.90', '41.77', '42.30', '41.50', '42.75'),
    ('200', '53.89', '54.10', '48.20', '49.20', '46.62', '48.45'),
    ('250', '66.33', '64.80', '58.19', '59.75', '54.98', '56.75'),
]

for row_idx, row_data in enumerate(t_data):
    for col_idx, val in enumerate(row_data):
        cell = table2.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph('Table 2: ')
p.runs[0].bold = True
p.runs[0].font.size = Pt(9)
run = p.add_run('V\u03c6 (cm\u00b3/mol) at 30 MPa. Experimental data from Hnedkovsky et al. (1996).')
run.font.size = Pt(9)
run.italic = True

doc.add_heading('4.3 CO\u2082 Density Predictions', level=2)

doc.add_paragraph(
    'Garcia (2001) reports a density increase of approximately 2.5% for CO\u2082 at x\u2082 = 0.05 '
    'and 25\u00b0C. Our model gives +2.69%, consistent with Garcia\'s result. The density '
    'increase is nearly linear with mole fraction, as Garcia observed.'
)

doc.add_heading('4.4 Physical Reasonableness', level=2)

doc.add_paragraph(
    'The following physical constraints are satisfied across all 8 gases:'
)

checks = [
    'V\u03c6 increases monotonically with temperature from 50\u2013250\u00b0C at 30 MPa',
    'CO\u2082 increases solution density (high MW relative to V\u03c6)',
    'H\u2082 decreases solution density most strongly (lowest MW)',
    'H\u2082S has near-neutral density effect (MW/V\u03c6 ratio close to water)',
    'V\u03c6 ordering follows molecular size: H\u2082 < N\u2082 \u2248 CO\u2082 \u2248 H\u2082S < CH\u2084 < C\u2082H\u2086 < C\u2083H\u2088 < n-C\u2084H\u2081\u2080',
    'Adding CH\u2084 to a CO\u2082 solution reduces the density increase (mixed gas test)',
]
for c in checks:
    doc.add_paragraph(c, style='List Bullet')

# ============================================================
doc.add_heading('5. Density Effects at Reservoir Conditions', level=1)

doc.add_paragraph(
    'Table 3 shows the temperature dependence of the CO\u2082 density effect at 30 MPa with '
    '2 mol% dissolved CO\u2082. The density increase diminishes at higher temperatures and '
    'reverses above ~240\u00b0C as V\u03c6 grows rapidly approaching the water critical point.'
)

table3 = doc.add_table(rows=7, cols=4)
table3.style = 'Light Shading Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['T (\u00b0C)', '\u03c1_water (kg/m\u00b3)', '\u03c1_solution (kg/m\u00b3)', '\u0394\u03c1/\u03c1']
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

t3_data = [
    ('25',  '1010.12', '1021.27', '+1.10%'),
    ('50',  '1000.67', '1012.14', '+1.15%'),
    ('100',  '971.82',  '980.14', '+0.86%'),
    ('150',  '932.86',  '937.97', '+0.55%'),
    ('200',  '884.62',  '885.93', '+0.15%'),
    ('250',  '825.56',  '821.98', '\u22120.43%'),
]

for row_idx, row_data in enumerate(t3_data):
    for col_idx, val in enumerate(row_data):
        cell = table3.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
p = doc.add_paragraph('Table 3: ')
p.runs[0].bold = True
p.runs[0].font.size = Pt(9)
run = p.add_run('CO\u2082 density correction vs temperature at P = 30 MPa, x\u2082 = 0.02.')
run.font.size = Pt(9)
run.italic = True

# ============================================================
doc.add_heading('6. Brine Support', level=1)

doc.add_paragraph(
    'Garcia (2001) showed that salinity effects on V\u03c6 are weak and within experimental '
    'uncertainty. Therefore, V\u03c6 values from the Plyasunov model (fitted to pure water data) '
    'are used regardless of salinity. The salinity enters only through the base solvent '
    'density \u03c1\u2081: for brines, \u03c1\u2081 is the CO\u2082-free brine density from the Spivey et al. '
    '(modified) correlation, rather than pure water density.'
)

doc.add_paragraph(
    'The density effect of dissolved gas is smaller in brine than in pure water '
    'because brine is denser, so the same mass of dissolved gas represents a smaller '
    'fractional change.'
)

# ============================================================
doc.add_heading('7. Scope and Limitations', level=1)

limits = [
    ('Temperature range', '25\u2013250\u00b0C (298\u2013523 K). The Plyasunov model is formally valid to higher '
     'temperatures, but V\u03c6 diverges approaching the water critical point (~374\u00b0C) and '
     'the Garcia mixing rule assumes dilute solutions.'),
    ('Pressure range', '0.1\u2013100 MPa. V\u03c6 has weak pressure dependence below 250\u00b0C.'),
    ('Concentration', 'Dilute solutions (x\u2082 < 0.05\u20130.10). The model uses infinite-dilution '
     'V\u03c6 values, appropriate for typical reservoir dissolved gas concentrations.'),
    ('Salinity', 'NaCl brines up to ~25 wt%. V\u03c6 values from pure water are used '
     '(salinity effect on V\u03c6 is within experimental uncertainty per Garcia 2001).'),
    ('Mixed gases', 'Ideal mixing of apparent molar volumes. Not validated against '
     'experimental mixed-gas density data (which is scarce to non-existent).'),
    ('Gases covered', 'CO\u2082, CH\u2084, C\u2082H\u2086, C\u2083H\u2088, n-C\u2084H\u2081\u2080, H\u2082S, H\u2082, N\u2082. '
     'Extension to other gases would require Plyasunov coefficients (available for '
     'O\u2082, CO, noble gases, and several polar molecules).'),
]

for label, text in limits:
    p = doc.add_paragraph()
    run = p.add_run(label + ': ')
    run.bold = True
    p.add_run(text)

# ============================================================
doc.add_heading('8. Water Properties', level=1)

doc.add_paragraph(
    'The Plyasunov model requires pure water density and isothermal compressibility '
    'at each (T, P) point. Two approaches are available:'
)

doc.add_paragraph(
    'IAPWS-95 (via iapws Python library): The full scientific formulation. Highest '
    'accuracy. Suitable for Python-based tools (pyResToolbox).',
    style='List Bullet'
)

doc.add_paragraph(
    'IAPWS-IF97 Region 1 (from scratch): The industrial formulation for compressed '
    'liquid water. A single 34-term polynomial in reduced pressure and temperature, '
    'requiring only its first and second derivatives. Approximately 130 lines of code '
    'with no external dependencies. Suitable for Dart/Flutter (ResToolbox3) or any '
    'environment without access to IAPWS libraries.',
    style='List Bullet'
)

doc.add_paragraph(
    'IF97 agrees with IAPWS-95 to within 0.014 kg/m\u00b3 on density and 0.2% on '
    'compressibility over the tested range (25\u2013300\u00b0C, 0.1\u2013100 MPa). '
    'Both are negligible relative to the Plyasunov model uncertainties. All 28 '
    'validation checks pass with either backend.'
)

doc.add_paragraph(
    'Note: the Spivey brine density correlation (used for the base brine density) '
    'should be retained as-is. Its salt correction equations depend on internal '
    'intermediate values (reference-pressure density, compressibility parameters) '
    'that cannot be replaced with IAPWS. IAPWS is used only for the Plyasunov '
    'model\u2019s pure water inputs. The two do not conflict.'
)

# ============================================================
doc.add_heading('9. References', level=1)

refs = [
    'Garcia, J.E. (2001). "Density of Aqueous Solutions of CO\u2082." LBNL-49023. DOI: 10.2172/790022.',
    'Hnedkovsky, L., Wood, R.H. & Majer, V. (1996). "Volumes of aqueous solutions of CH\u2084, CO\u2082, '
    'H\u2082S, and NH\u2083 at temperatures from 298.15 K to 705 K and pressures to 35 MPa." '
    'J. Chem. Thermodynamics 28, 125-142.',
    'Plyasunov, A.V. & Korzhinskaya, V.S. (2020). "Thermodynamic properties of aqueous solutions. '
    'Part II: CO\u2082, C\u2082H\u2084, C\u2082H\u2086, C\u2083H\u2088, n-C\u2084H\u2081\u2080, i-C\u2084H\u2081\u2080." '
    'Fluid Phase Equilibria 521, 112690.',
    'Plyasunov, A.V. & Korzhinskaya, V.S. (2021a). "Thermodynamic properties of aqueous solutions. '
    'Part III: H\u2082S and polar solutes." Fluid Phase Equilibria, 112872.',
    'Plyasunov, A.V. & Korzhinskaya, V.S. (2021b). "Thermodynamic properties of aqueous solutions. '
    'Part IV: Re-parametrization for H\u2082, N\u2082, O\u2082, CO, CH\u2084." '
    'Fluid Phase Equilibria 536, 112982.',
    'Wagner, W. et al. (2000). "The IAPWS Industrial Formulation 1997 for the Thermodynamic '
    'Properties of Water and Steam." ASME J. Eng. Gas Turbines Power, 122(1), 150-182.',
]

for r in refs:
    doc.add_paragraph(r, style='List Number')

# Save
doc.save('/home/mark/projects/garcia_extension/Multi-Gas_Density_Corrections_Technical_Note.docx')
print("Technical note saved.")
